"""
ContractIQ Export
Generate reports in PDF, Word, and HTML formats.
"""

import logging
from typing import Optional, Dict, List
from datetime import datetime
import io
import base64

from .models import ContractAnalysis, RiskLevel

logger = logging.getLogger(__name__)


class ContractExporter:
    """
    Export contract analysis to various formats.
    Supports PDF, Word (DOCX), HTML, and JSON.
    """

    def __init__(self):
        """Initialize exporter."""
        pass

    def export_to_html(
        self,
        analysis: ContractAnalysis,
        include_sections: Optional[List[str]] = None
    ) -> str:
        """
        Export analysis to HTML format.

        Args:
            analysis: ContractAnalysis to export
            include_sections: Specific sections to include (None = all)

        Returns:
            HTML string
        """
        sections = include_sections or [
            "summary", "parties", "dates", "payment",
            "termination", "liability", "risks", "recommendations"
        ]

        html = self._get_html_header(analysis)

        # Executive Summary
        if "summary" in sections:
            html += self._render_summary_section(analysis)

        # Parties
        if "parties" in sections and analysis.parties:
            html += self._render_parties_section(analysis)

        # Key Dates
        if "dates" in sections and (analysis.effective_date or analysis.key_dates):
            html += self._render_dates_section(analysis)

        # Payment Terms
        if "payment" in sections and analysis.payment_terms:
            html += self._render_payment_section(analysis)

        # Termination
        if "termination" in sections and analysis.termination_clauses:
            html += self._render_termination_section(analysis)

        # Liability
        if "liability" in sections and analysis.liability_clauses:
            html += self._render_liability_section(analysis)

        # Risk Assessment
        if "risks" in sections:
            html += self._render_risk_section(analysis)

        # Recommendations
        if "recommendations" in sections and analysis.recommendations:
            html += self._render_recommendations_section(analysis)

        html += self._get_html_footer()

        return html

    def export_to_pdf(
        self,
        analysis: ContractAnalysis,
        include_sections: Optional[List[str]] = None
    ) -> bytes:
        """
        Export analysis to PDF format.

        Requires: weasyprint or reportlab

        Args:
            analysis: ContractAnalysis to export
            include_sections: Specific sections to include

        Returns:
            PDF as bytes
        """
        try:
            # Try WeasyPrint first
            from weasyprint import HTML
            html_content = self.export_to_html(analysis, include_sections)
            pdf_bytes = HTML(string=html_content).write_pdf()
            return pdf_bytes
        except ImportError:
            logger.warning("WeasyPrint not installed, trying reportlab")
            return self._export_pdf_reportlab(analysis, include_sections)
        except Exception as e:
            logger.error(f"PDF export failed: {e}")
            raise

    def _export_pdf_reportlab(
        self,
        analysis: ContractAnalysis,
        include_sections: Optional[List[str]] = None
    ) -> bytes:
        """Export to PDF using reportlab."""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib import colors

            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []

            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30
            )
            story.append(Paragraph("Contract Analysis Report", title_style))
            story.append(Spacer(1, 12))

            # Contract info
            story.append(Paragraph(
                f"<b>Contract:</b> {analysis.contract_title or 'Untitled'}",
                styles['Normal']
            ))
            story.append(Paragraph(
                f"<b>Type:</b> {analysis.contract_type.value.title()}",
                styles['Normal']
            ))
            story.append(Paragraph(
                f"<b>Analyzed:</b> {analysis.analyzed_at}",
                styles['Normal']
            ))
            story.append(Spacer(1, 24))

            # Risk Score
            risk_color = {
                RiskLevel.LOW: colors.green,
                RiskLevel.MEDIUM: colors.orange,
                RiskLevel.HIGH: colors.red,
                RiskLevel.CRITICAL: colors.darkred
            }.get(analysis.risk_level, colors.gray)

            story.append(Paragraph(
                f"<b>Risk Score:</b> {analysis.risk_score}/100 ({analysis.risk_level.value.upper()})",
                styles['Heading2']
            ))
            story.append(Spacer(1, 12))

            # Summary
            story.append(Paragraph("Executive Summary", styles['Heading2']))
            story.append(Paragraph(analysis.summary, styles['Normal']))
            story.append(Spacer(1, 12))

            # Key Points
            if analysis.key_points:
                story.append(Paragraph("Key Points", styles['Heading2']))
                for point in analysis.key_points:
                    story.append(Paragraph(f"• {point}", styles['Normal']))
                story.append(Spacer(1, 12))

            # Parties
            if analysis.parties:
                story.append(Paragraph("Parties", styles['Heading2']))
                for party in analysis.parties:
                    story.append(Paragraph(
                        f"<b>{party.role}:</b> {party.name}",
                        styles['Normal']
                    ))
                story.append(Spacer(1, 12))

            # Risk Flags
            if analysis.risk_flags:
                story.append(Paragraph("Risk Flags", styles['Heading2']))
                for flag in analysis.risk_flags:
                    severity = flag.severity.value.upper()
                    story.append(Paragraph(
                        f"<b>[{severity}]</b> {flag.description}",
                        styles['Normal']
                    ))
                story.append(Spacer(1, 12))

            # Recommendations
            if analysis.recommendations:
                story.append(Paragraph("Recommendations", styles['Heading2']))
                for rec in analysis.recommendations:
                    story.append(Paragraph(f"• {rec}", styles['Normal']))

            doc.build(story)
            return buffer.getvalue()

        except ImportError:
            logger.error("reportlab not installed")
            raise ImportError("Please install reportlab or weasyprint for PDF export")

    def export_to_docx(
        self,
        analysis: ContractAnalysis,
        include_sections: Optional[List[str]] = None
    ) -> bytes:
        """
        Export analysis to Word (DOCX) format.

        Requires: python-docx

        Args:
            analysis: ContractAnalysis to export
            include_sections: Specific sections to include

        Returns:
            DOCX as bytes
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Title
            title = doc.add_heading("Contract Analysis Report", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Contract Info
            doc.add_paragraph()
            info = doc.add_paragraph()
            info.add_run("Contract: ").bold = True
            info.add_run(analysis.contract_title or "Untitled")

            info = doc.add_paragraph()
            info.add_run("Type: ").bold = True
            info.add_run(analysis.contract_type.value.title())

            info = doc.add_paragraph()
            info.add_run("Analyzed: ").bold = True
            info.add_run(analysis.analyzed_at)

            # Risk Score
            doc.add_heading("Risk Assessment", level=1)
            risk_para = doc.add_paragraph()
            risk_para.add_run("Risk Score: ").bold = True
            risk_run = risk_para.add_run(
                f"{analysis.risk_score}/100 ({analysis.risk_level.value.upper()})"
            )
            risk_color = {
                RiskLevel.LOW: RGBColor(0, 128, 0),
                RiskLevel.MEDIUM: RGBColor(255, 165, 0),
                RiskLevel.HIGH: RGBColor(255, 0, 0),
                RiskLevel.CRITICAL: RGBColor(139, 0, 0)
            }.get(analysis.risk_level, RGBColor(128, 128, 128))
            risk_run.font.color.rgb = risk_color

            # Executive Summary
            doc.add_heading("Executive Summary", level=1)
            doc.add_paragraph(analysis.summary)

            # Key Points
            if analysis.key_points:
                doc.add_heading("Key Points", level=1)
                for point in analysis.key_points:
                    doc.add_paragraph(point, style='List Bullet')

            # Parties
            if analysis.parties:
                doc.add_heading("Parties", level=1)
                for party in analysis.parties:
                    para = doc.add_paragraph()
                    para.add_run(f"{party.role}: ").bold = True
                    para.add_run(party.name)

            # Key Dates
            if analysis.effective_date or analysis.expiration_date:
                doc.add_heading("Key Dates", level=1)
                if analysis.effective_date:
                    para = doc.add_paragraph()
                    para.add_run("Effective Date: ").bold = True
                    para.add_run(analysis.effective_date)
                if analysis.expiration_date:
                    para = doc.add_paragraph()
                    para.add_run("Expiration Date: ").bold = True
                    para.add_run(analysis.expiration_date)

            # Payment Terms
            if analysis.payment_terms:
                doc.add_heading("Payment Terms", level=1)
                for payment in analysis.payment_terms:
                    doc.add_paragraph(
                        f"• {payment.amount} {payment.currency} - {payment.frequency or 'One-time'}",
                        style='List Bullet'
                    )

            # Risk Flags
            if analysis.risk_flags:
                doc.add_heading("Risk Flags", level=1)
                for flag in analysis.risk_flags:
                    para = doc.add_paragraph()
                    severity_run = para.add_run(f"[{flag.severity.value.upper()}] ")
                    severity_run.bold = True
                    if flag.severity in [RiskLevel.HIGH, RiskLevel.CRITICAL]:
                        severity_run.font.color.rgb = RGBColor(255, 0, 0)
                    para.add_run(flag.description)

            # Recommendations
            if analysis.recommendations:
                doc.add_heading("Recommendations", level=1)
                for rec in analysis.recommendations:
                    doc.add_paragraph(rec, style='List Bullet')

            # Save to bytes
            buffer = io.BytesIO()
            doc.save(buffer)
            return buffer.getvalue()

        except ImportError:
            logger.error("python-docx not installed")
            raise ImportError("Please install python-docx for Word export")

    def export_to_json(self, analysis: ContractAnalysis) -> str:
        """Export analysis to JSON format."""
        return analysis.model_dump_json(indent=2)

    def get_email_content(
        self,
        analysis: ContractAnalysis,
        recipient_name: Optional[str] = None
    ) -> Dict[str, str]:
        """
        Generate email content for contract analysis.

        Returns:
            Dict with 'subject', 'body_text', 'body_html'
        """
        subject = f"Contract Analysis: {analysis.contract_title or 'Untitled'} - Risk: {analysis.risk_level.value.upper()}"

        # Plain text body
        body_text = f"""
Contract Analysis Report
========================

Contract: {analysis.contract_title or 'Untitled'}
Type: {analysis.contract_type.value.title()}
Risk Score: {analysis.risk_score}/100 ({analysis.risk_level.value.upper()})

Executive Summary:
{analysis.summary}

Key Points:
"""
        for point in analysis.key_points:
            body_text += f"• {point}\n"

        if analysis.risk_flags:
            body_text += "\nRisk Flags:\n"
            for flag in analysis.risk_flags[:5]:
                body_text += f"• [{flag.severity.value.upper()}] {flag.description}\n"

        if analysis.recommendations:
            body_text += "\nRecommendations:\n"
            for rec in analysis.recommendations[:5]:
                body_text += f"• {rec}\n"

        body_text += "\n\n---\nGenerated by ContractIQ"

        # HTML body
        body_html = self.export_to_html(analysis, ["summary", "risks", "recommendations"])

        return {
            "subject": subject,
            "body_text": body_text,
            "body_html": body_html
        }

    # =========================================================================
    # HTML RENDERING HELPERS
    # =========================================================================

    def _get_html_header(self, analysis: ContractAnalysis) -> str:
        """Generate HTML header."""
        risk_color = {
            RiskLevel.LOW: "#10B981",
            RiskLevel.MEDIUM: "#F59E0B",
            RiskLevel.HIGH: "#EF4444",
            RiskLevel.CRITICAL: "#7F1D1D"
        }.get(analysis.risk_level, "#6B7280")

        return f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Contract Analysis Report</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 40px 20px;
            color: #1F2937;
        }}
        h1 {{ color: #111827; border-bottom: 2px solid #E5E7EB; padding-bottom: 10px; }}
        h2 {{ color: #374151; margin-top: 30px; }}
        .risk-badge {{
            display: inline-block;
            padding: 8px 16px;
            border-radius: 9999px;
            color: white;
            font-weight: bold;
            background-color: {risk_color};
        }}
        .section {{ margin-bottom: 30px; }}
        .party {{ padding: 10px; background: #F3F4F6; margin: 5px 0; border-radius: 4px; }}
        .risk-flag {{ padding: 10px; margin: 5px 0; border-left: 4px solid; }}
        .risk-low {{ border-color: #10B981; background: #ECFDF5; }}
        .risk-medium {{ border-color: #F59E0B; background: #FFFBEB; }}
        .risk-high {{ border-color: #EF4444; background: #FEF2F2; }}
        .risk-critical {{ border-color: #7F1D1D; background: #FEE2E2; }}
        .recommendation {{ padding: 10px; background: #EFF6FF; margin: 5px 0; border-radius: 4px; }}
        table {{ width: 100%; border-collapse: collapse; margin: 10px 0; }}
        th, td {{ padding: 10px; text-align: left; border-bottom: 1px solid #E5E7EB; }}
        th {{ background: #F9FAFB; }}
        .footer {{ margin-top: 40px; padding-top: 20px; border-top: 1px solid #E5E7EB; color: #6B7280; font-size: 12px; }}
    </style>
</head>
<body>
    <h1>Contract Analysis Report</h1>
    
    <div class="section">
        <p><strong>Contract:</strong> {analysis.contract_title or 'Untitled'}</p>
        <p><strong>Type:</strong> {analysis.contract_type.value.title()}</p>
        <p><strong>Analyzed:</strong> {analysis.analyzed_at}</p>
        <p><strong>Risk Score:</strong> <span class="risk-badge">{analysis.risk_score}/100 - {analysis.risk_level.value.upper()}</span></p>
    </div>
"""

    def _get_html_footer(self) -> str:
        """Generate HTML footer."""
        return f"""
    <div class="footer">
        <p>Generated by ContractIQ • {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}</p>
    </div>
</body>
</html>
"""

    def _render_summary_section(self, analysis: ContractAnalysis) -> str:
        """Render summary section."""
        html = '<div class="section">'
        html += '<h2>Executive Summary</h2>'
        html += f'<p>{analysis.summary}</p>'

        if analysis.key_points:
            html += '<h3>Key Points</h3><ul>'
            for point in analysis.key_points:
                html += f'<li>{point}</li>'
            html += '</ul>'

        html += '</div>'
        return html

    def _render_parties_section(self, analysis: ContractAnalysis) -> str:
        """Render parties section."""
        html = '<div class="section">'
        html += '<h2>Parties</h2>'

        for party in analysis.parties:
            html += f'''
            <div class="party">
                <strong>{party.role}:</strong> {party.name}
                {f'<br><small>{party.address}</small>' if party.address else ''}
            </div>
            '''

        html += '</div>'
        return html

    def _render_dates_section(self, analysis: ContractAnalysis) -> str:
        """Render dates section."""
        html = '<div class="section">'
        html += '<h2>Key Dates</h2>'
        html += '<table>'
        html += '<tr><th>Date Type</th><th>Date</th></tr>'

        if analysis.effective_date:
            html += f'<tr><td>Effective Date</td><td>{analysis.effective_date}</td></tr>'
        if analysis.expiration_date:
            html += f'<tr><td>Expiration Date</td><td>{analysis.expiration_date}</td></tr>'

        for date in analysis.key_dates:
            html += f'<tr><td>{date.date_type}</td><td>{date.date_text}</td></tr>'

        html += '</table></div>'
        return html

    def _render_payment_section(self, analysis: ContractAnalysis) -> str:
        """Render payment section."""
        html = '<div class="section">'
        html += '<h2>Payment Terms</h2>'
        html += '<table>'
        html += '<tr><th>Amount</th><th>Frequency</th><th>Due</th></tr>'

        for payment in analysis.payment_terms:
            html += f'''
            <tr>
                <td>{payment.amount} {payment.currency}</td>
                <td>{payment.frequency or 'One-time'}</td>
                <td>{payment.due_date or 'N/A'}</td>
            </tr>
            '''

        if analysis.total_value:
            html += f'<tr><td colspan="3"><strong>Total Value: {analysis.total_value}</strong></td></tr>'

        html += '</table></div>'
        return html

    def _render_termination_section(self, analysis: ContractAnalysis) -> str:
        """Render termination section."""
        html = '<div class="section">'
        html += '<h2>Termination Clauses</h2>'

        for clause in analysis.termination_clauses:
            html += f'''
            <div class="party">
                <strong>{clause.termination_type.replace('_', ' ').title()}</strong><br>
                Notice Period: {clause.notice_period or 'Not specified'}<br>
                {f'Penalties: {clause.penalties}' if clause.penalties else ''}
            </div>
            '''

        html += '</div>'
        return html

    def _render_liability_section(self, analysis: ContractAnalysis) -> str:
        """Render liability section."""
        html = '<div class="section">'
        html += '<h2>Liability & Indemnification</h2>'

        for clause in analysis.liability_clauses:
            html += f'''
            <div class="party">
                <strong>{clause.liability_type.title()}</strong><br>
                {f'Cap: {clause.cap_amount}' if clause.cap_amount else 'No cap specified'}<br>
                {f'Excluded: {", ".join(clause.excluded_damages)}' if clause.excluded_damages else ''}
            </div>
            '''

        html += '</div>'
        return html

    def _render_risk_section(self, analysis: ContractAnalysis) -> str:
        """Render risk section."""
        html = '<div class="section">'
        html += '<h2>Risk Assessment</h2>'

        if analysis.risk_flags:
            for flag in analysis.risk_flags:
                risk_class = f"risk-{flag.severity.value}"
                html += f'''
                <div class="risk-flag {risk_class}">
                    <strong>[{flag.severity.value.upper()}] {flag.risk_type}</strong><br>
                    {flag.description}
                    {f'<br><em>Recommendation: {flag.recommendation}</em>' if flag.recommendation else ''}
                </div>
                '''
        else:
            html += '<p>No significant risks identified.</p>'

        if analysis.missing_clauses:
            html += '<h3>Missing Standard Clauses</h3><ul>'
            for clause in analysis.missing_clauses:
                html += f'<li>{clause}</li>'
            html += '</ul>'

        html += '</div>'
        return html

    def _render_recommendations_section(self, analysis: ContractAnalysis) -> str:
        """Render recommendations section."""
        html = '<div class="section">'
        html += '<h2>Recommendations</h2>'

        for rec in analysis.recommendations:
            html += f'<div class="recommendation">{rec}</div>'

        html += '</div>'
        return html

